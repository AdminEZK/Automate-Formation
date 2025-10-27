#!/usr/bin/env python3
"""
Script pour ajouter les variables directement dans les templates Word
Ce script ouvre chaque template et ajoute les variables aux bons endroits
"""

import os
from docx import Document
from docx.shared import Pt
import shutil

TEMPLATES_DIR = "Dossier exemple"
BACKUP_DIR = os.path.join(TEMPLATES_DIR, "backup_originaux")

def backup_template(template_name):
    """Crée une sauvegarde du template original"""
    if not os.path.exists(BACKUP_DIR):
        os.makedirs(BACKUP_DIR)
    
    source = os.path.join(TEMPLATES_DIR, template_name)
    backup = os.path.join(BACKUP_DIR, template_name)
    
    if os.path.exists(source) and not os.path.exists(backup):
        shutil.copy2(source, backup)
        print(f"💾 Sauvegarde créée : {template_name}")
        return True
    return False

def replace_text_in_runs(paragraph, old_text, new_text):
    """Remplace du texte dans un paragraphe en préservant le style"""
    full_text = ''.join(run.text for run in paragraph.runs)
    
    if old_text in full_text:
        # Trouver le premier run qui contient le texte
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True
    return False

def add_variables_to_convention():
    """Ajoute les variables au template de convention"""
    template_name = "Modèle de convention simplifiée de formation 2020.docx"
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    
    if not os.path.exists(template_path):
        print(f"❌ Template non trouvé : {template_name}")
        return
    
    print(f"\n📄 Traitement : {template_name}")
    backup_template(template_name)
    
    try:
        doc = Document(template_path)
        
        # Liste des remplacements à faire
        replacements = {
            # Ces valeurs seront remplacées par les variables
            # Format: "texte_à_chercher": "{{variable}}"
            
            # Vous devrez adapter ces valeurs selon votre template
            # Exemple:
            # "Aladé Conseil": "{{organisme_raison_sociale}}",
            # "SIRET : 12345678900012": "SIRET : {{organisme_siret}}",
        }
        
        # Parcourir tous les paragraphes
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                replace_text_in_runs(paragraph, old_text, new_text)
        
        # Parcourir tous les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            replace_text_in_runs(paragraph, old_text, new_text)
        
        # Sauvegarder
        doc.save(template_path)
        print(f"✅ Variables ajoutées avec succès")
        
    except Exception as e:
        print(f"❌ Erreur : {str(e)}")

def create_template_from_scratch(template_name, content_structure):
    """Crée un nouveau template avec les variables déjà en place"""
    template_path = os.path.join(TEMPLATES_DIR, f"Template_{template_name}")
    
    doc = Document()
    
    # Ajouter le contenu avec les variables
    for item in content_structure:
        if item['type'] == 'heading':
            doc.add_heading(item['text'], level=item.get('level', 1))
        elif item['type'] == 'paragraph':
            doc.add_paragraph(item['text'])
        elif item['type'] == 'table':
            table = doc.add_table(rows=len(item['data']), cols=len(item['data'][0]))
            for i, row_data in enumerate(item['data']):
                for j, cell_data in enumerate(row_data):
                    table.rows[i].cells[j].text = cell_data
    
    doc.save(template_path)
    print(f"✅ Template créé : {template_path}")

def create_convention_template():
    """Crée un template de convention avec toutes les variables"""
    content = [
        {'type': 'heading', 'text': 'CONVENTION DE FORMATION PROFESSIONNELLE', 'level': 1},
        {'type': 'paragraph', 'text': '(Articles L. 6353-1 et L. 6353-2 du Code du travail)'},
        {'type': 'heading', 'text': 'ENTRE LES SOUSSIGNÉS :', 'level': 2},
        
        {'type': 'heading', 'text': "L'ORGANISME DE FORMATION", 'level': 3},
        {'type': 'paragraph', 'text': '{{organisme_raison_sociale}}'},
        {'type': 'paragraph', 'text': 'SIRET : {{organisme_siret}}'},
        {'type': 'paragraph', 'text': "N° de déclaration d'activité : {{organisme_nda}}"},
        {'type': 'paragraph', 'text': 'Adresse : {{organisme_adresse}}, {{organisme_code_postal}} {{organisme_ville}}'},
        {'type': 'paragraph', 'text': 'Représenté par : {{organisme_representant_nom}}, {{organisme_representant_fonction}}'},
        {'type': 'paragraph', 'text': 'Email : {{organisme_email}}'},
        {'type': 'paragraph', 'text': 'Téléphone : {{organisme_telephone}}'},
        
        {'type': 'paragraph', 'text': "D'UNE PART,"},
        {'type': 'paragraph', 'text': ''},
        {'type': 'paragraph', 'text': 'ET'},
        {'type': 'paragraph', 'text': ''},
        
        {'type': 'heading', 'text': 'LE CLIENT', 'level': 3},
        {'type': 'paragraph', 'text': '{{entreprise_nom}}'},
        {'type': 'paragraph', 'text': 'SIRET : {{entreprise_siret}}'},
        {'type': 'paragraph', 'text': 'Adresse : {{entreprise_adresse}}, {{entreprise_code_postal}} {{entreprise_ville}}'},
        {'type': 'paragraph', 'text': 'Email : {{entreprise_email}}'},
        {'type': 'paragraph', 'text': 'Téléphone : {{entreprise_telephone}}'},
        
        {'type': 'paragraph', 'text': "D'AUTRE PART,"},
        {'type': 'paragraph', 'text': ''},
        
        {'type': 'heading', 'text': 'IL A ÉTÉ CONVENU CE QUI SUIT :', 'level': 2},
        
        {'type': 'heading', 'text': "ARTICLE 1 - OBJET DE LA CONVENTION", 'level': 3},
        {'type': 'paragraph', 'text': 'La présente convention a pour objet la réalisation d\'une action de formation professionnelle continue intitulée "{{formation_titre}}" conformément aux dispositions des articles L. 6353-1 et suivants du Code du travail.'},
        
        {'type': 'heading', 'text': "ARTICLE 2 - NATURE ET CARACTÉRISTIQUES DE L'ACTION", 'level': 3},
        {'type': 'table', 'data': [
            ['Intitulé', '{{formation_titre}}'],
            ['Durée', '{{formation_duree}} heures'],
            ['Dates', 'Du {{date_debut}} au {{date_fin}}'],
            ['Horaires', 'De {{horaire_debut}} à {{horaire_fin}}'],
            ['Lieu', '{{lieu}}'],
            ['Nombre de participants', '{{nombre_participants}}'],
        ]},
        
        {'type': 'heading', 'text': 'ARTICLE 3 - MODALITÉS FINANCIÈRES', 'level': 3},
        {'type': 'paragraph', 'text': 'Le coût total de la formation s\'élève à {{prix_total_ht}}, soit {{prix_total_ttc}}.'},
        
        {'type': 'heading', 'text': 'ARTICLE 4 - DÉLAI DE RÉTRACTATION', 'level': 3},
        {'type': 'paragraph', 'text': 'Conformément à l\'article L. 6353-5 du Code du travail, le client dispose d\'un délai de 10 jours à compter de la signature de la présente convention pour se rétracter.'},
        
        {'type': 'paragraph', 'text': ''},
        {'type': 'paragraph', 'text': 'Fait en deux exemplaires originaux,'},
        {'type': 'paragraph', 'text': 'À {{organisme_ville}}, le {{date_aujourd_hui}}'},
        {'type': 'paragraph', 'text': ''},
        
        {'type': 'table', 'data': [
            ["Pour l'organisme de formation", 'Pour le client'],
            ['{{organisme_representant_nom}}', ''],
            ['Signature :', 'Signature :'],
        ]},
    ]
    
    create_template_from_scratch('Convention.docx', content)

def create_convocation_template():
    """Crée un template de convocation avec toutes les variables"""
    content = [
        {'type': 'heading', 'text': '{{organisme_nom}}', 'level': 1},
        {'type': 'paragraph', 'text': ''},
        {'type': 'heading', 'text': 'CONVOCATION À UNE FORMATION', 'level': 1},
        {'type': 'paragraph', 'text': ''},
        {'type': 'paragraph', 'text': '{{participant_prenom}} {{participant_nom}}'},
        {'type': 'paragraph', 'text': '{{participant_fonction}}'},
        {'type': 'paragraph', 'text': '{{entreprise_nom}}'},
        {'type': 'paragraph', 'text': ''},
        {'type': 'paragraph', 'text': 'Madame, Monsieur,'},
        {'type': 'paragraph', 'text': ''},
        {'type': 'paragraph', 'text': 'Nous avons le plaisir de vous convoquer à la formation suivante :'},
        {'type': 'paragraph', 'text': ''},
        {'type': 'table', 'data': [
            ['Formation', '{{formation_titre}}'],
            ['Dates', 'Du {{date_debut}} au {{date_fin}}'],
            ['Horaires', 'De {{horaire_debut}} à {{horaire_fin}}'],
            ['Durée', '{{formation_duree}} heures'],
            ['Lieu', '{{lieu}}'],
            ['Formateur', '{{formateur_nom_complet}}'],
        ]},
        {'type': 'paragraph', 'text': ''},
        {'type': 'paragraph', 'text': 'Merci de vous présenter 15 minutes avant le début de la formation avec une pièce d\'identité.'},
        {'type': 'paragraph', 'text': ''},
        {'type': 'paragraph', 'text': 'Pour toute question, n\'hésitez pas à nous contacter :'},
        {'type': 'paragraph', 'text': 'Email : {{organisme_email}}'},
        {'type': 'paragraph', 'text': 'Téléphone : {{organisme_telephone}}'},
        {'type': 'paragraph', 'text': ''},
        {'type': 'paragraph', 'text': 'Nous vous souhaitons une excellente formation.'},
        {'type': 'paragraph', 'text': ''},
        {'type': 'paragraph', 'text': 'Cordialement,'},
        {'type': 'paragraph', 'text': '{{organisme_representant_nom}}'},
        {'type': 'paragraph', 'text': '{{organisme_representant_fonction}}'},
    ]
    
    create_template_from_scratch('Convocation.docx', content)

def main():
    print("🚀 Ajout des variables dans les templates Word")
    print("=" * 60)
    print()
    print("⚠️  ATTENTION : Ce script va modifier vos templates Word.")
    print("   Une sauvegarde sera créée dans 'backup_originaux/'")
    print()
    
    choice = input("Voulez-vous continuer ? (o/n) : ")
    
    if choice.lower() != 'o':
        print("❌ Opération annulée")
        return
    
    print("\n📋 Options disponibles :")
    print("1. Créer de nouveaux templates avec variables (recommandé)")
    print("2. Modifier les templates existants (expérimental)")
    print()
    
    option = input("Choisissez une option (1 ou 2) : ")
    
    if option == '1':
        print("\n✨ Création de nouveaux templates...")
        create_convention_template()
        create_convocation_template()
        print("\n✅ Templates créés dans le dossier 'Dossier exemple/'")
        print("   Noms : Template_Convention.docx, Template_Convocation.docx")
        print("   Vous pouvez les utiliser comme référence pour modifier vos templates existants.")
        
    elif option == '2':
        print("\n⚠️  Cette option est expérimentale.")
        print("   Il est recommandé de modifier manuellement vos templates.")
        print("   Consultez PREPARATION-TEMPLATES.md pour les instructions.")
        
    else:
        print("❌ Option invalide")
    
    print("\n" + "=" * 60)
    print("✅ Terminé !")

if __name__ == '__main__':
    main()
