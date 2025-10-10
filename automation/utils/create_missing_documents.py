"""
Création des 4 documents manquants pour le workflow Qualiopi
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_evaluation_opco():
    """
    Tâche 10 : Évaluation OPCO (Financeur)
    Basé sur les exigences Qualiopi pour les financeurs
    """
    doc = Document()
    
    # En-tête
    title = doc.add_heading('QUESTIONNAIRE D\'ÉVALUATION FINANCEUR (OPCO)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informations formation
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Organisme de formation : ').bold = True
    p.add_run('[NOM_ORGANISME]')
    
    p = doc.add_paragraph()
    p.add_run('N° de déclaration d\'activité : ').bold = True
    p.add_run('[NDA]')
    
    p = doc.add_paragraph()
    p.add_run('Intitulé de la formation : ').bold = True
    p.add_run('[TITRE_FORMATION]')
    
    p = doc.add_paragraph()
    p.add_run('Entreprise bénéficiaire : ').bold = True
    p.add_run('[NOM_ENTREPRISE]')
    
    p = doc.add_paragraph()
    p.add_run('Dates de réalisation : ').bold = True
    p.add_run('Du [DATE_DEBUT] au [DATE_FIN]')
    
    p = doc.add_paragraph()
    p.add_run('Nombre de participants : ').bold = True
    p.add_run('[NOMBRE_PARTICIPANTS]')
    
    doc.add_paragraph()
    doc.add_heading('ÉVALUATION DE LA PRESTATION', 2)
    
    # Questions conformes Qualiopi
    questions = [
        {
            'question': '1. Les objectifs de la formation ont-ils été atteints ?',
            'options': '☐ Totalement  ☐ Partiellement  ☐ Pas du tout'
        },
        {
            'question': '2. La formation a-t-elle répondu aux besoins identifiés de l\'entreprise ?',
            'options': '☐ Oui  ☐ Partiellement  ☐ Non'
        },
        {
            'question': '3. Les compétences visées ont-elles été acquises par les participants ?',
            'options': '☐ Oui  ☐ Partiellement  ☐ Non  ☐ Ne sait pas'
        },
        {
            'question': '4. Les moyens pédagogiques et techniques étaient-ils adaptés ?',
            'options': '☐ Oui  ☐ Partiellement  ☐ Non'
        },
        {
            'question': '5. Le déroulement de la formation s\'est-il effectué conformément au programme ?',
            'options': '☐ Oui  ☐ Non'
        },
        {
            'question': '6. Les documents fournis (convention, programme, attestations) étaient-ils conformes ?',
            'options': '☐ Oui  ☐ Non'
        },
        {
            'question': '7. Recommanderiez-vous cet organisme de formation ?',
            'options': '☐ Oui  ☐ Non'
        }
    ]
    
    for q in questions:
        doc.add_paragraph()
        p = doc.add_paragraph(q['question'])
        p.runs[0].bold = True
        doc.add_paragraph(q['options'])
    
    # Commentaires
    doc.add_paragraph()
    doc.add_heading('OBSERVATIONS ET COMMENTAIRES', 2)
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Date : ').bold = True
    p.add_run('_' * 20)
    
    p = doc.add_paragraph()
    p.add_run('Nom et fonction du représentant OPCO : ').bold = True
    p.add_run('_' * 40)
    
    p = doc.add_paragraph()
    p.add_run('Signature : ').bold = True
    
    # Sauvegarder
    output_path = 'Dossier exemple/Modèle Évaluation OPCO.docx'
    doc.save(output_path)
    print(f'✅ Document créé : {output_path}')
    return output_path


def create_deroule_pedagogique():
    """
    Tâche 12 : Déroulé pédagogique
    Document interne pour l'organisme de formation (exigence Qualiopi)
    """
    doc = Document()
    
    # En-tête
    title = doc.add_heading('DÉROULÉ PÉDAGOGIQUE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Document interne - Organisme de formation')
    doc.add_paragraph()
    
    # Informations formation
    doc.add_heading('INFORMATIONS GÉNÉRALES', 2)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    
    rows_data = [
        ('Intitulé de la formation', '[TITRE_FORMATION]'),
        ('Formateur', '[NOM_FORMATEUR]'),
        ('Dates', 'Du [DATE_DEBUT] au [DATE_FIN]'),
        ('Durée totale', '[DUREE] heures'),
        ('Lieu', '[LIEU]'),
        ('Entreprise', '[NOM_ENTREPRISE]'),
        ('Nombre de participants', '[NOMBRE_PARTICIPANTS]'),
        ('Format', '[FORMAT]')  # présentiel, distanciel, mixte
    ]
    
    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
    
    # Objectifs
    doc.add_paragraph()
    doc.add_heading('OBJECTIFS PÉDAGOGIQUES', 2)
    doc.add_paragraph('[OBJECTIFS_FORMATION]')
    
    # Public visé
    doc.add_paragraph()
    doc.add_heading('PUBLIC VISÉ', 2)
    doc.add_paragraph('[PUBLIC_VISE]')
    
    # Prérequis
    doc.add_paragraph()
    doc.add_heading('PRÉREQUIS', 2)
    doc.add_paragraph('[PREREQUIS]')
    
    # Déroulé détaillé
    doc.add_paragraph()
    doc.add_heading('DÉROULÉ DÉTAILLÉ PAR SÉQUENCE', 2)
    
    # Tableau déroulé
    table_deroule = doc.add_table(rows=1, cols=5)
    table_deroule.style = 'Light Grid Accent 1'
    
    # En-têtes
    headers = ['Séquence', 'Durée', 'Contenu', 'Méthodes', 'Supports']
    for i, header in enumerate(headers):
        cell = table_deroule.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Exemples de séquences (placeholders)
    sequences = [
        ('Accueil et présentation', '0h30', '[Contenu]', '[Méthodes]', '[Supports]'),
        ('Séquence 1', '[Durée]', '[Contenu]', '[Méthodes]', '[Supports]'),
        ('Pause', '0h15', 'Pause', '-', '-'),
        ('Séquence 2', '[Durée]', '[Contenu]', '[Méthodes]', '[Supports]'),
        ('Déjeuner', '1h00', 'Pause déjeuner', '-', '-'),
        ('Séquence 3', '[Durée]', '[Contenu]', '[Méthodes]', '[Supports]'),
        ('Évaluation et clôture', '0h30', '[Contenu]', '[Méthodes]', '[Supports]')
    ]
    
    for seq in sequences:
        row = table_deroule.add_row()
        for i, value in enumerate(seq):
            row.cells[i].text = value
    
    # Moyens pédagogiques
    doc.add_paragraph()
    doc.add_heading('MOYENS PÉDAGOGIQUES ET TECHNIQUES', 2)
    doc.add_paragraph('[MOYENS_PEDAGOGIQUES]')
    
    # Modalités d'évaluation
    doc.add_paragraph()
    doc.add_heading('MODALITÉS D\'ÉVALUATION', 2)
    doc.add_paragraph('[MODALITES_EVALUATION]')
    
    # Notes
    doc.add_paragraph()
    doc.add_heading('NOTES ET OBSERVATIONS', 2)
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    
    output_path = 'Dossier exemple/Modèle Déroulé Pédagogique.docx'
    doc.save(output_path)
    print(f'✅ Document créé : {output_path}')
    return output_path


def create_questionnaire_formateur():
    """
    Tâche 13 : Questionnaire formateur (Bilan après formation)
    Conforme aux exigences Qualiopi
    """
    doc = Document()
    
    # En-tête
    title = doc.add_heading('QUESTIONNAIRE D\'ÉVALUATION FORMATEUR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Bilan de la formation')
    doc.add_paragraph()
    
    # Informations
    p = doc.add_paragraph()
    p.add_run('Formateur : ').bold = True
    p.add_run('[NOM_FORMATEUR]')
    
    p = doc.add_paragraph()
    p.add_run('Formation : ').bold = True
    p.add_run('[TITRE_FORMATION]')
    
    p = doc.add_paragraph()
    p.add_run('Dates : ').bold = True
    p.add_run('Du [DATE_DEBUT] au [DATE_FIN]')
    
    p = doc.add_paragraph()
    p.add_run('Entreprise : ').bold = True
    p.add_run('[NOM_ENTREPRISE]')
    
    doc.add_paragraph()
    doc.add_heading('DÉROULEMENT DE LA FORMATION', 2)
    
    questions = [
        {
            'question': '1. Le déroulement de la formation s\'est-il effectué conformément au programme prévu ?',
            'options': '☐ Oui  ☐ Non  ☐ Partiellement'
        },
        {
            'question': 'Si non ou partiellement, précisez :',
            'options': '_' * 80
        },
        {
            'question': '2. Les objectifs pédagogiques ont-ils été atteints ?',
            'options': '☐ Totalement  ☐ Partiellement  ☐ Pas du tout'
        },
        {
            'question': '3. Le niveau du groupe était-il homogène ?',
            'options': '☐ Oui  ☐ Non'
        },
        {
            'question': '4. Les participants étaient-ils impliqués et motivés ?',
            'options': '☐ Très impliqués  ☐ Moyennement  ☐ Peu impliqués'
        },
        {
            'question': '5. Les moyens pédagogiques et techniques mis à disposition étaient-ils adaptés ?',
            'options': '☐ Oui  ☐ Non  ☐ Partiellement'
        },
        {
            'question': '6. Les conditions matérielles (salle, équipements) étaient-elles satisfaisantes ?',
            'options': '☐ Oui  ☐ Non'
        },
        {
            'question': 'Si non, précisez :',
            'options': '_' * 80
        },
        {
            'question': '7. Avez-vous rencontré des difficultés particulières ?',
            'options': '☐ Oui  ☐ Non'
        },
        {
            'question': 'Si oui, lesquelles :',
            'options': '_' * 80
        }
    ]
    
    for q in questions:
        doc.add_paragraph()
        p = doc.add_paragraph(q['question'])
        if not q['question'].startswith('Si'):
            p.runs[0].bold = True
        doc.add_paragraph(q['options'])
    
    # Observations
    doc.add_paragraph()
    doc.add_heading('OBSERVATIONS ET SUGGESTIONS D\'AMÉLIORATION', 2)
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    doc.add_paragraph('_' * 80)
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Date : ').bold = True
    p.add_run('_' * 20)
    
    p = doc.add_paragraph()
    p.add_run('Signature du formateur : ').bold = True
    
    output_path = 'Dossier exemple/Modèle Questionnaire Formateur.docx'
    doc.save(output_path)
    print(f'✅ Document créé : {output_path}')
    return output_path


def create_contrat_formateur():
    """
    Tâche 15 : Contrat de sous-traitance formateur
    Conforme aux exigences Qualiopi et droit du travail
    """
    doc = Document()
    
    # En-tête
    title = doc.add_heading('CONTRAT DE SOUS-TRAITANCE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_sub = doc.add_heading('Prestation de formation professionnelle', 2)
    title_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Parties
    doc.add_heading('ENTRE LES SOUSSIGNÉS :', 2)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('L\'organisme de formation :').bold = True
    
    doc.add_paragraph('[NOM_ORGANISME_FORMATION]')
    doc.add_paragraph('SIRET : [SIRET_ORGANISME]')
    doc.add_paragraph('N° de déclaration d\'activité : [NDA]')
    doc.add_paragraph('Adresse : [ADRESSE_ORGANISME]')
    doc.add_paragraph('Représenté par : [REPRESENTANT_LEGAL], [FONCTION]')
    
    doc.add_paragraph()
    p = doc.add_paragraph('Ci-après dénommé « le Donneur d\'ordre »')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('D\'UNE PART,').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('ET :').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('[NOM_FORMATEUR] [PRENOM_FORMATEUR]')
    doc.add_paragraph('N° SIRET/SIREN : [SIRET_FORMATEUR]')
    doc.add_paragraph('Adresse : [ADRESSE_FORMATEUR]')
    doc.add_paragraph('Email : [EMAIL_FORMATEUR]')
    doc.add_paragraph('Téléphone : [TELEPHONE_FORMATEUR]')
    
    doc.add_paragraph()
    p = doc.add_paragraph('Ci-après dénommé « le Formateur »')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('D\'AUTRE PART,').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Articles
    doc.add_page_break()
    
    doc.add_heading('ARTICLE 1 - OBJET DU CONTRAT', 2)
    doc.add_paragraph(
        'Le présent contrat a pour objet de définir les conditions dans lesquelles le Formateur '
        'intervient pour le compte du Donneur d\'ordre dans le cadre d\'une action de formation professionnelle.'
    )
    
    doc.add_heading('ARTICLE 2 - MISSION', 2)
    doc.add_paragraph('Le Formateur s\'engage à réaliser la prestation suivante :')
    doc.add_paragraph()
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    mission_data = [
        ('Intitulé de la formation', '[TITRE_FORMATION]'),
        ('Thématique', '[THEMATIQUE]'),
        ('Durée', '[DUREE] heures'),
        ('Dates d\'intervention', 'Du [DATE_DEBUT] au [DATE_FIN]'),
        ('Lieu', '[LIEU]'),
        ('Entreprise bénéficiaire', '[NOM_ENTREPRISE]'),
        ('Nombre de participants', '[NOMBRE_PARTICIPANTS]')
    ]
    
    for i, (label, value) in enumerate(mission_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
    
    doc.add_heading('ARTICLE 3 - OBLIGATIONS DU FORMATEUR', 2)
    doc.add_paragraph('Le Formateur s\'engage à :')
    obligations = [
        'Réaliser la prestation conformément au programme de formation fourni',
        'Respecter les horaires convenus',
        'Fournir les supports pédagogiques nécessaires',
        'Remplir et signer les feuilles d\'émargement',
        'Respecter les exigences du référentiel national qualité (Qualiopi)',
        'Informer immédiatement le Donneur d\'ordre de toute difficulté rencontrée',
        'Respecter la confidentialité des informations de l\'entreprise bénéficiaire',
        'Fournir un bilan de formation dans les 7 jours suivant la fin de la prestation'
    ]
    for obligation in obligations:
        doc.add_paragraph(f'• {obligation}')
    
    doc.add_heading('ARTICLE 4 - OBLIGATIONS DU DONNEUR D\'ORDRE', 2)
    doc.add_paragraph('Le Donneur d\'ordre s\'engage à :')
    obligations_do = [
        'Fournir au Formateur toutes les informations nécessaires à la réalisation de la mission',
        'Mettre à disposition les moyens matériels et techniques nécessaires',
        'Régler les honoraires selon les modalités prévues à l\'article 5',
        'Informer le Formateur de toute modification du planning'
    ]
    for obligation in obligations_do:
        doc.add_paragraph(f'• {obligation}')
    
    doc.add_heading('ARTICLE 5 - RÉMUNÉRATION', 2)
    doc.add_paragraph('La rémunération du Formateur est fixée comme suit :')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Taux horaire : ').bold = True
    p.add_run('[TAUX_HORAIRE] € HT / heure')
    
    p = doc.add_paragraph()
    p.add_run('Montant total : ').bold = True
    p.add_run('[MONTANT_TOTAL] € HT')
    
    doc.add_paragraph()
    doc.add_paragraph('Modalités de paiement : Règlement à 30 jours sur présentation de facture')
    
    doc.add_heading('ARTICLE 6 - EXIGENCES QUALIOPI', 2)
    doc.add_paragraph(
        'Le Formateur reconnaît être informé que le Donneur d\'ordre est certifié Qualiopi '
        'et s\'engage à respecter les exigences du référentiel national qualité, notamment :'
    )
    exigences = [
        'Respect du programme de formation',
        'Adaptation des méthodes pédagogiques au public',
        'Évaluation des acquis des participants',
        'Fourniture des documents requis (émargement, bilan)',
        'Respect des délais et des engagements'
    ]
    for exigence in exigences:
        doc.add_paragraph(f'• {exigence}')
    
    doc.add_heading('ARTICLE 7 - ASSURANCES', 2)
    doc.add_paragraph(
        'Le Formateur déclare être titulaire d\'une assurance responsabilité civile professionnelle '
        'couvrant les dommages qui pourraient être causés dans le cadre de son activité.'
    )
    
    doc.add_heading('ARTICLE 8 - RÉSILIATION', 2)
    doc.add_paragraph(
        'En cas de manquement grave de l\'une des parties à ses obligations, le présent contrat '
        'pourra être résilié de plein droit par l\'autre partie, 8 jours après l\'envoi d\'une mise '
        'en demeure restée sans effet.'
    )
    
    doc.add_heading('ARTICLE 9 - LITIGES', 2)
    doc.add_paragraph(
        'Tout litige relatif à l\'interprétation ou à l\'exécution du présent contrat sera soumis '
        'aux tribunaux compétents.'
    )
    
    # Signatures
    doc.add_page_break()
    doc.add_heading('SIGNATURES', 2)
    
    doc.add_paragraph()
    doc.add_paragraph('Fait en deux exemplaires originaux')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('À : ').bold = True
    p.add_run('_' * 30)
    
    p = doc.add_paragraph()
    p.add_run('Le : ').bold = True
    p.add_run('_' * 30)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Tableau signatures
    table_sign = doc.add_table(rows=4, cols=2)
    
    table_sign.rows[0].cells[0].text = 'Pour le Donneur d\'ordre'
    table_sign.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table_sign.rows[0].cells[1].text = 'Pour le Formateur'
    table_sign.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    table_sign.rows[1].cells[0].text = '[REPRESENTANT_LEGAL]'
    table_sign.rows[1].cells[1].text = '[NOM_FORMATEUR] [PRENOM_FORMATEUR]'
    
    table_sign.rows[2].cells[0].text = '[FONCTION]'
    table_sign.rows[2].cells[1].text = ''
    
    table_sign.rows[3].cells[0].text = '\n\n\nSignature :'
    table_sign.rows[3].cells[1].text = '\n\n\nSignature :'
    
    output_path = 'Dossier exemple/Modèle Contrat Formateur.docx'
    doc.save(output_path)
    print(f'✅ Document créé : {output_path}')
    return output_path


if __name__ == '__main__':
    print('=' * 80)
    print('CRÉATION DES DOCUMENTS MANQUANTS POUR QUALIOPI')
    print('=' * 80)
    print()
    
    create_evaluation_opco()
    create_deroule_pedagogique()
    create_questionnaire_formateur()
    create_contrat_formateur()
    
    print()
    print('=' * 80)
    print('✅ TOUS LES DOCUMENTS ONT ÉTÉ CRÉÉS AVEC SUCCÈS !')
    print('=' * 80)
    print()
    print('📁 Emplacement : Dossier exemple/')
    print()
    print('Documents créés :')
    print('1. Modèle Évaluation OPCO.docx')
    print('2. Modèle Déroulé Pédagogique.docx')
    print('3. Modèle Questionnaire Formateur.docx')
    print('4. Modèle Contrat Formateur.docx')
    print()
    print('⚠️  NOTES IMPORTANTES :')
    print('- Tous les placeholders sont entre crochets [PLACEHOLDER]')
    print('- Les documents respectent les exigences Qualiopi')
    print('- Vérifiez et adaptez selon vos besoins spécifiques')
    print('- Le contrat formateur doit être validé par un juriste')
