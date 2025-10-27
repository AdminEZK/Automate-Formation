#!/usr/bin/env python3
"""
Script pour modifier automatiquement les templates Word
Remplace les valeurs fixes par des variables dynamiques
"""

import os
from docx import Document
import shutil
from datetime import datetime

TEMPLATES_DIR = "Dossier exemple"
BACKUP_DIR = os.path.join(TEMPLATES_DIR, "backup_originaux")

def backup_templates():
    """Crée une sauvegarde de tous les templates"""
    if not os.path.exists(BACKUP_DIR):
        os.makedirs(BACKUP_DIR)
        print(f"📁 Dossier de sauvegarde créé : {BACKUP_DIR}\n")
    
    templates = [
        "Modele Convocation formation 2020.docx",
        "Modèle Proposition de Formation_2.docx",
        "Modèle de convention simplifiée de formation 2020.docx",
        "Modèle Programme de formation 2020.docx",
        "Modèle Certificat de réalisation.docx",
    ]
    
    for template in templates:
        source = os.path.join(TEMPLATES_DIR, template)
        backup = os.path.join(BACKUP_DIR, template)
        
        if os.path.exists(source) and not os.path.exists(backup):
            shutil.copy2(source, backup)
            print(f"💾 Sauvegarde : {template}")

def replace_in_paragraph(paragraph, replacements):
    """Remplace le texte dans un paragraphe en préservant le style"""
    for old_text, new_text in replacements.items():
        if old_text in paragraph.text:
            # Remplacer dans chaque run pour préserver le style
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

def replace_in_table(table, replacements):
    """Remplace le texte dans un tableau"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)

def replace_in_document(doc, replacements):
    """Remplace le texte dans tout le document"""
    # Paragraphes
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    
    # Tableaux
    for table in doc.tables:
        replace_in_table(table, replacements)
    
    # En-têtes
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for table in section.header.tables:
            replace_in_table(table, replacements)
    
    # Pieds de page
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for table in section.footer.tables:
            replace_in_table(table, replacements)

def modifier_convocation():
    """Modifie le template de convocation"""
    template_name = "Modele Convocation formation 2020.docx"
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    
    if not os.path.exists(template_path):
        print(f"❌ Template non trouvé : {template_name}")
        return
    
    print(f"\n📄 Modification : {template_name}")
    
    try:
        doc = Document(template_path)
        
        replacements = {
            # Lieu et date
            'Saint-Malo': '{{organisme_ville}}',
            'le 27 septembre 2021': 'le {{date_aujourd_hui}}',
            
            # Participant
            '« NOM / PRENOM »': '{{participant_nom_complet}}',
            
            # Tableau - Intitulé formation
            'Renouvellement CACES R386 CATÉGORIE 3B': '{{formation_titre}}',
            
            # Durée
            '14': '{{formation_duree}}',
            
            # Dates (dans le tableau)
            '15/11/2019': '{{date_debut}}',
            '22/11/2019': '{{date_fin}}',
            
            # Horaires
            '8H30 -16H30': '{{horaire_debut}} - {{horaire_fin}}',
            
            # Lieu de formation
            'Entreprise MAHEY\n5, Impasse Grand Jardin\n35400 Saint Malo': '{{lieu}}',
        }
        
        replace_in_document(doc, replacements)
        
        doc.save(template_path)
        print(f"✅ Convocation modifiée avec succès")
        
    except Exception as e:
        print(f"❌ Erreur : {str(e)}")

def modifier_proposition():
    """Modifie le template de proposition de formation"""
    template_name = "Modèle Proposition de Formation_2.docx"
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    
    if not os.path.exists(template_path):
        print(f"❌ Template non trouvé : {template_name}")
        return
    
    print(f"\n📄 Modification : {template_name}")
    
    try:
        doc = Document(template_path)
        
        replacements = {
            # Titre
            'Nom de la formation': '{{formation_titre}}',
            'Nom du client': '{{entreprise_nom}}',
            
            # Section I - Contexte
            'Société FAST SUD': '{{entreprise_nom}}',
            '2 rue des Peupliers': '{{entreprise_adresse}}',
            '35 380 PLELAN-LE-GRAND': '{{entreprise_code_postal}} {{entreprise_ville}}',
            'Spécialisée dans le secteur d\'activité de la restauration de type rapide': '{{entreprise_description}}',
            
            # Section II - Demande
            'XXXXXXXXle nom de la formationXXXXXXXXXXXX': '{{formation_titre}}',
            
            # Section III - Proposition
            '5 septembre 2022': '{{date_proposition}}',
            '8 septembre 2022': '{{date_entrevue}}',
            
            # Section IV - Objectifs
            'Mise en conformité réglementaire sur la conduite d\'engin ……': '{{formation_objectifs}}',
            
            # Section V - Programme
            'Rôle et compétences du Maître d\'apprentissage en entreprise': '{{formation_programme}}',
            
            # Section VI - Public
            'Public : XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX': 'Public : {{formation_public_vise}}',
            'Prérequis : XXXXXXXXXXXXXXXXXXXXXXXXXX': 'Prérequis : {{formation_prerequis}}',
            
            # Section VII - Organisation
            'Présentiel en INTRA': '{{formation_modalite}}',
            'gilles.pichonnet@free.fr': '{{organisme_email}}',
            'Gilles PICHONNET': '{{organisme_representant_nom}}',
            
            # Durée
            '14 heures, soit 2 jours.': '{{formation_duree}} heures',
            
            # Dates et lieu
            'A définir d\'un commun accord.': 'Du {{date_debut}} au {{date_fin}}',
            'En intra dans vos locaux': '{{lieu}}',
            
            # Section IX - Prix
            '2 400€': '{{prix_total_ht}}',
            '480€': '{{tva}}',
            '2 880€': '{{prix_total_ttc}}',
            
            # Signature
            'Saint-Malo': '{{organisme_ville}}',
            'le 5 septembre 2024': 'le {{date_aujourd_hui}}',
            'Monsieur XXXXXXXXXX': '{{entreprise_representant_nom}}',
            'Titre : XXXXXXXXXX': '{{entreprise_representant_fonction}}',
            'Monsieur Gilles PICHONNET': '{{organisme_representant_nom}}',
            'Dirigeant Fondateur de PICHFORMATION': '{{organisme_representant_fonction}}',
        }
        
        replace_in_document(doc, replacements)
        
        doc.save(template_path)
        print(f"✅ Proposition modifiée avec succès")
        
    except Exception as e:
        print(f"❌ Erreur : {str(e)}")

def modifier_convention():
    """Modifie le template de convention"""
    template_name = "Modèle de convention simplifiée de formation 2020.docx"
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    
    if not os.path.exists(template_path):
        print(f"❌ Template non trouvé : {template_name}")
        return
    
    print(f"\n📄 Modification : {template_name}")
    
    try:
        doc = Document(template_path)
        
        # Remplacements génériques pour la convention
        # Ces valeurs seront adaptées selon le contenu réel de votre template
        replacements = {
            # À adapter selon votre template
            'PICHFORMATION': '{{organisme_raison_sociale}}',
            'Gilles PICHONNET': '{{organisme_representant_nom}}',
            'Saint-Malo': '{{organisme_ville}}',
        }
        
        replace_in_document(doc, replacements)
        
        doc.save(template_path)
        print(f"✅ Convention modifiée avec succès")
        
    except Exception as e:
        print(f"❌ Erreur : {str(e)}")

def modifier_programme():
    """Modifie le template de programme"""
    template_name = "Modèle Programme de formation 2020.docx"
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    
    if not os.path.exists(template_path):
        print(f"❌ Template non trouvé : {template_name}")
        return
    
    print(f"\n📄 Modification : {template_name}")
    
    try:
        doc = Document(template_path)
        
        replacements = {
            'PICHFORMATION': '{{organisme_nom}}',
            'Gilles PICHONNET': '{{organisme_representant_nom}}',
            'gilles.pichonnet@free.fr': '{{organisme_email}}',
            'Saint-Malo': '{{organisme_ville}}',
        }
        
        replace_in_document(doc, replacements)
        
        doc.save(template_path)
        print(f"✅ Programme modifié avec succès")
        
    except Exception as e:
        print(f"❌ Erreur : {str(e)}")

def modifier_certificat():
    """Modifie le template de certificat"""
    template_name = "Modèle Certificat de réalisation.docx"
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    
    if not os.path.exists(template_path):
        print(f"❌ Template non trouvé : {template_name}")
        return
    
    print(f"\n📄 Modification : {template_name}")
    
    try:
        doc = Document(template_path)
        
        replacements = {
            'PICHFORMATION': '{{organisme_raison_sociale}}',
            'Gilles PICHONNET': '{{organisme_representant_nom}}',
            'Saint-Malo': '{{organisme_ville}}',
        }
        
        replace_in_document(doc, replacements)
        
        doc.save(template_path)
        print(f"✅ Certificat modifié avec succès")
        
    except Exception as e:
        print(f"❌ Erreur : {str(e)}")

def main():
    print("🚀 Modification automatique des templates Word")
    print("=" * 60)
    print()
    print("⚠️  ATTENTION : Ce script va modifier vos templates Word.")
    print("   Une sauvegarde sera créée dans 'backup_originaux/'")
    print()
    
    choice = input("Voulez-vous continuer ? (o/n) : ")
    
    if choice.lower() != 'o':
        print("❌ Opération annulée")
        return
    
    print("\n📦 Création des sauvegardes...")
    backup_templates()
    
    print("\n🔧 Modification des templates...")
    
    # Modifier chaque template
    modifier_convocation()
    modifier_proposition()
    modifier_convention()
    modifier_programme()
    modifier_certificat()
    
    print("\n" + "=" * 60)
    print("✅ Modification terminée !")
    print(f"📁 Les originaux sont sauvegardés dans : {BACKUP_DIR}")
    print("\n💡 Prochaines étapes :")
    print("   1. Ouvrez les templates modifiés pour vérifier")
    print("   2. Testez la génération avec : python test-document-generation.py")
    print("   3. Ajustez manuellement si nécessaire")

if __name__ == '__main__':
    main()
