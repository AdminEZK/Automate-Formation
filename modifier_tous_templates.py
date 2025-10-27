#!/usr/bin/env python3
"""
Script pour modifier TOUS les templates Word du dossier exemple
Ajoute les variables appropriées dans chaque document
"""

import os
from docx import Document
import shutil
from datetime import datetime

TEMPLATES_DIR = "Dossier exemple"
BACKUP_DIR = os.path.join(TEMPLATES_DIR, "backup_originaux")

def backup_all_templates():
    """Crée une sauvegarde de TOUS les templates"""
    if not os.path.exists(BACKUP_DIR):
        os.makedirs(BACKUP_DIR)
        print(f"📁 Dossier de sauvegarde créé : {BACKUP_DIR}\n")
    
    templates = [
        "Modele Convocation formation 2020.docx",
        "Modele bulletin d'inscription formation INTER 2020.docx",
        "Modèle Certificat de réalisation.docx",
        "Modèle Évaluation de la satisfaction du client.docx",
        "Modèle Feuille d émargement entreprise.docx",
        "Modèle Feuille d émargement individuelle.docx",
        "Modèle Grille de MAJ des compétences.docx",
        "Modèle Programme de formation 2020.docx",
        "Modèle Proposition de Formation_2.docx",
        "Modèle Questionnaire préalable à la formation.docx",
        "Modèle Traitement des réclamations majeures.docx",
        "Modèle de convention simplifiée de formation 2020.docx",
        "Modèle évaluation à chaud.docx",
        "Modèle évaluation à froid.docx",
        "Modèle Contrat Formateur.docx",
        "Modèle Déroulé Pédagogique.docx",
        "Modèle Questionnaire Formateur.docx",
        "Modèle Évaluation OPCO.docx",
        "Règlement Intérieur.docx",
    ]
    
    count = 0
    for template in templates:
        source = os.path.join(TEMPLATES_DIR, template)
        backup = os.path.join(BACKUP_DIR, template)
        
        if os.path.exists(source):
            if not os.path.exists(backup):
                shutil.copy2(source, backup)
                print(f"💾 {template}")
                count += 1
    
    print(f"\n✅ {count} templates sauvegardés\n")

def replace_in_paragraph(paragraph, replacements):
    """Remplace le texte dans un paragraphe"""
    for old_text, new_text in replacements.items():
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

def replace_in_table(table, replacements):
    """Remplace le texte dans un tableau"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)

def replace_in_document(doc, replacements):
    """Remplace le texte dans tout le document"""
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    
    for table in doc.tables:
        replace_in_table(table, replacements)
    
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for table in section.header.tables:
            replace_in_table(table, replacements)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for table in section.footer.tables:
            replace_in_table(table, replacements)

def get_common_replacements():
    """Retourne les remplacements communs à tous les documents"""
    return {
        # Organisme (à adapter selon vos valeurs)
        'PICHFORMATION': '{{organisme_nom}}',
        'Gilles PICHONNET': '{{organisme_representant_nom}}',
        'Dirigeant Fondateur': '{{organisme_representant_fonction}}',
        'gilles.pichonnet@free.fr': '{{organisme_email}}',
        # Saint-Malo sera remplacé spécifiquement dans chaque template
        '35400 Saint-Malo': '{{organisme_code_postal}} {{organisme_ville}}',
        '35400 Saint Malo': '{{organisme_code_postal}} {{organisme_ville}}',
    }

def modifier_template(template_name, specific_replacements=None):
    """Modifie un template avec les remplacements spécifiques"""
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    
    if not os.path.exists(template_path):
        print(f"⚠️  Non trouvé : {template_name}")
        return False
    
    print(f"📄 {template_name}")
    
    try:
        doc = Document(template_path)
        
        # Remplacements communs
        replacements = get_common_replacements()
        
        # Ajouter les remplacements spécifiques
        if specific_replacements:
            replacements.update(specific_replacements)
        
        replace_in_document(doc, replacements)
        doc.save(template_path)
        print(f"   ✅ Modifié\n")
        return True
        
    except Exception as e:
        print(f"   ❌ Erreur : {str(e)}\n")
        return False

def main():
    print("🚀 Modification de TOUS les templates Word")
    print("=" * 60)
    print()
    
    # Sauvegarde
    print("📦 Création des sauvegardes...")
    backup_all_templates()
    
    print("🔧 Modification des templates...\n")
    
    success_count = 0
    
    # 1. CONVOCATION
    if modifier_template("Modele Convocation formation 2020.docx", {
        '« NOM / PRENOM »': '{{participant_prenom}} {{participant_nom}}',
        'NOM / PRENOM': '{{participant_prenom}} {{participant_nom}}',
        'le 27 septembre 2021': 'le {{date_aujourd_hui}}',
        'Renouvellement CACES R386 CATÉGORIE 3B': '{{formation_titre}}',
        '14': '{{formation_duree}}',
        '15/11/2019': '{{date_debut}}',
        '22/11/2019': '{{date_fin}}',
        '8H30 -16H30': '{{horaire_debut}} - {{horaire_fin}}',
        '8H30': '{{horaire_debut}}',
        '16H30': '{{horaire_fin}}',
        'Entreprise MAHEY': '{{lieu}}',
        '5, Impasse Grand Jardin': '',
        'Saint Malo': '',
        '35400 Saint Malo': '{{organisme_code_postal}} {{organisme_ville}}',
        '{{organisme_code_postal}}\nSaint Malo': '{{organisme_code_postal}} {{organisme_ville}}',
    }):
        success_count += 1
    
    # 2. PROPOSITION DE FORMATION
    if modifier_template("Modèle Proposition de Formation_2.docx", {
        'Nom de la formation': '{{formation_titre}}',
        'Nom du client': '{{entreprise_nom}}',
        'Société FAST SUD': '{{entreprise_nom}}',
        '2 rue des Peupliers': '{{entreprise_adresse}}',
        '35 380 PLELAN-LE-GRAND': '{{entreprise_code_postal}} {{entreprise_ville}}',
        'XXXXXXXXle nom de la formationXXXXXXXXXXXX': '{{formation_titre}}',
        '5 septembre 2022': '{{date_proposition}}',
        '8 septembre 2022': '{{date_entrevue}}',
        'Présentiel en INTRA': '{{formation_modalite}}',
        '14 heures, soit 2 jours.': '{{formation_duree}} heures',
        'A définir d\'un commun accord.': 'Du {{date_debut}} au {{date_fin}}',
        'En intra dans vos locaux': '{{lieu}}',
        '2 400€': '{{prix_total_ht}}',
        '480€': '{{tva}}',
        '2 880€': '{{prix_total_ttc}}',
        'le 5 septembre 2024': 'le {{date_aujourd_hui}}',
        'Monsieur XXXXXXXXXX': '{{entreprise_representant_nom}}',
        'Titre : XXXXXXXXXX': '{{entreprise_representant_fonction}}',
    }):
        success_count += 1
    
    # 3. CONVENTION DE FORMATION
    if modifier_template("Modèle de convention simplifiée de formation 2020.docx", {
        # Variables spécifiques à la convention
    }):
        success_count += 1
    
    # 4. PROGRAMME DE FORMATION
    if modifier_template("Modèle Programme de formation 2020.docx", {
        # Variables spécifiques au programme
    }):
        success_count += 1
    
    # 5. CERTIFICAT DE RÉALISATION
    if modifier_template("Modèle Certificat de réalisation.docx", {
        # Variables participant
    }):
        success_count += 1
    
    # 6. FEUILLE ÉMARGEMENT ENTREPRISE
    if modifier_template("Modèle Feuille d émargement entreprise.docx", {
        # Variables émargement
    }):
        success_count += 1
    
    # 7. FEUILLE ÉMARGEMENT INDIVIDUELLE
    if modifier_template("Modèle Feuille d émargement individuelle.docx", {
        # Variables émargement individuel
    }):
        success_count += 1
    
    # 8. QUESTIONNAIRE PRÉALABLE
    if modifier_template("Modèle Questionnaire préalable à la formation.docx", {
        # Variables questionnaire
    }):
        success_count += 1
    
    # 9. ÉVALUATION À CHAUD
    if modifier_template("Modèle évaluation à chaud.docx", {
        # Variables évaluation
    }):
        success_count += 1
    
    # 10. ÉVALUATION À FROID
    if modifier_template("Modèle évaluation à froid.docx", {
        # Variables évaluation
    }):
        success_count += 1
    
    # 11. ÉVALUATION SATISFACTION CLIENT
    if modifier_template("Modèle Évaluation de la satisfaction du client.docx", {
        # Variables évaluation client
    }):
        success_count += 1
    
    # 12. RÈGLEMENT INTÉRIEUR
    if modifier_template("Règlement Intérieur.docx", {
        # Variables règlement
    }):
        success_count += 1
    
    # 13. BULLETIN D'INSCRIPTION
    if modifier_template("Modele bulletin d'inscription formation INTER 2020.docx", {
        'ALADE CONSEILS': '{{organisme_nom}}',
        '22 rue de la Clarisse': '{{organisme_adresse}}',
        'Saint Malo': '{{organisme_ville}}',
        '06 85 66 43 70': '{{organisme_telephone}}',
        'contact@aladeconseils.com': '{{organisme_email}}',
        'Cliquez ou appuyez ici pour entrer du texte.': '',
        '800 € HT': '{{prix_unitaire_ht}}',
        '960 € TTC': '{{prix_unitaire_ttc}}',
    }):
        success_count += 1
    
    # 14. GRILLE MAJ COMPÉTENCES
    if modifier_template("Modèle Grille de MAJ des compétences.docx", {
        # Variables grille
    }):
        success_count += 1
    
    # 15. CONTRAT FORMATEUR
    if modifier_template("Modèle Contrat Formateur.docx", {
        # Variables contrat formateur
    }):
        success_count += 1
    
    # 16. DÉROULÉ PÉDAGOGIQUE
    if modifier_template("Modèle Déroulé Pédagogique.docx", {
        # Variables déroulé
    }):
        success_count += 1
    
    # 17. QUESTIONNAIRE FORMATEUR
    if modifier_template("Modèle Questionnaire Formateur.docx", {
        # Variables questionnaire formateur
    }):
        success_count += 1
    
    # 18. ÉVALUATION OPCO
    if modifier_template("Modèle Évaluation OPCO.docx", {
        # Variables OPCO
    }):
        success_count += 1
    
    # 19. TRAITEMENT RÉCLAMATIONS
    if modifier_template("Modèle Traitement des réclamations majeures.docx", {
        # Variables réclamations
    }):
        success_count += 1
    
    print("=" * 60)
    print(f"✅ {success_count} templates modifiés avec succès !")
    print(f"📁 Originaux sauvegardés dans : {BACKUP_DIR}")
    print("\n💡 Prochaines étapes :")
    print("   1. Ouvrez les templates pour vérifier")
    print("   2. Testez : python test-document-generation.py")
    print("   3. Ajustez manuellement si besoin")

if __name__ == '__main__':
    main()
