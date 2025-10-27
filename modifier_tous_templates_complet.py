#!/usr/bin/env python3
"""
Script complet pour modifier TOUS les templates avec les variables appropriées
"""

import os
from docx import Document
import shutil

TEMPLATES_DIR = "Dossier exemple"
BACKUP_DIR = os.path.join(TEMPLATES_DIR, "backup_originaux")

def clean_paragraph(paragraph, new_text):
    """Nettoie un paragraphe et ajoute le nouveau texte"""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    if new_text:
        paragraph.add_run(new_text)

def replace_in_doc(doc, replacements):
    """Remplace le texte dans tout le document"""
    # Paragraphes
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                clean_paragraph(paragraph, paragraph.text.replace(old_text, new_text))
    
    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            clean_paragraph(paragraph, paragraph.text.replace(old_text, new_text))
    
    # En-têtes et pieds de page
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    clean_paragraph(paragraph, paragraph.text.replace(old_text, new_text))
        for paragraph in section.footer.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    clean_paragraph(paragraph, paragraph.text.replace(old_text, new_text))

def get_common_replacements():
    """Remplacements communs à tous les documents"""
    return {
        'PICHFORMATION': '{{organisme_nom}}',
        'Gilles PICHONNET': '{{organisme_representant_nom}}',
        'gilles.pichonnet@free.fr': '{{organisme_email}}',
        '06 85 66 43 70': '{{organisme_telephone}}',
        'Dirigeant Fondateur': '{{organisme_representant_fonction}}',
        'Saint-Malo': '{{organisme_ville}}',
        'Saint Malo': '{{organisme_ville}}',
    }

print("🚀 Modification complète de TOUS les templates")
print("=" * 70)

# 1. CERTIFICAT DE RÉALISATION
print("\n📄 1. Certificat de réalisation...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Certificat de réalisation.docx"))
    replacements = get_common_replacements()
    replacements.update({
        # À adapter selon le contenu réel du certificat
        'Monsieur': '{{participant_prenom}}',
        'DUPONT': '{{participant_nom}}',
    })
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Certificat de réalisation.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 2. CONVOCATION (déjà fait)
print("\n📄 2. Convocation...")
print("   ✅ Déjà modifié")

# 3. PROPOSITION (déjà fait)
print("\n📄 3. Proposition de formation...")
print("   ✅ Déjà modifié")

# 4. CONVENTION
print("\n📄 4. Convention de formation...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle de convention simplifiée de formation 2020.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle de convention simplifiée de formation 2020.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 5. PROGRAMME
print("\n📄 5. Programme de formation...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Programme de formation 2020.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Programme de formation 2020.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 6. FEUILLE ÉMARGEMENT ENTREPRISE
print("\n📄 6. Feuille d'émargement entreprise...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Feuille d émargement entreprise.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Feuille d émargement entreprise.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 7. FEUILLE ÉMARGEMENT INDIVIDUELLE
print("\n📄 7. Feuille d'émargement individuelle...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Feuille d émargement individuelle.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Feuille d émargement individuelle.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 8. QUESTIONNAIRE PRÉALABLE
print("\n📄 8. Questionnaire préalable...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Questionnaire préalable à la formation.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Questionnaire préalable à la formation.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 9. ÉVALUATION À CHAUD
print("\n📄 9. Évaluation à chaud...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle évaluation à chaud.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle évaluation à chaud.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 10. ÉVALUATION À FROID
print("\n📄 10. Évaluation à froid...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle évaluation à froid.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle évaluation à froid.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 11. ÉVALUATION CLIENT
print("\n📄 11. Évaluation satisfaction client...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Évaluation de la satisfaction du client.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Évaluation de la satisfaction du client.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 12. RÈGLEMENT INTÉRIEUR
print("\n📄 12. Règlement intérieur...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Règlement Intérieur.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Règlement Intérieur.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 13. BULLETIN D'INSCRIPTION (déjà fait)
print("\n📄 13. Bulletin d'inscription...")
print("   ✅ Déjà modifié")

# 14. GRILLE COMPÉTENCES
print("\n📄 14. Grille MAJ compétences...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Grille de MAJ des compétences.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Grille de MAJ des compétences.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 15. CONTRAT FORMATEUR
print("\n📄 15. Contrat formateur...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Contrat Formateur.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Contrat Formateur.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 16. DÉROULÉ PÉDAGOGIQUE
print("\n📄 16. Déroulé pédagogique...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Déroulé Pédagogique.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Déroulé Pédagogique.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 17. QUESTIONNAIRE FORMATEUR
print("\n📄 17. Questionnaire formateur...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Questionnaire Formateur.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Questionnaire Formateur.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 18. ÉVALUATION OPCO
print("\n📄 18. Évaluation OPCO...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Évaluation OPCO.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Évaluation OPCO.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

# 19. TRAITEMENT RÉCLAMATIONS
print("\n📄 19. Traitement réclamations...")
try:
    doc = Document(os.path.join(BACKUP_DIR, "Modèle Traitement des réclamations majeures.docx"))
    replacements = get_common_replacements()
    replace_in_doc(doc, replacements)
    doc.save(os.path.join(TEMPLATES_DIR, "Modèle Traitement des réclamations majeures.docx"))
    print("   ✅ Modifié")
except Exception as e:
    print(f"   ❌ Erreur : {e}")

print("\n" + "=" * 70)
print("✅ Tous les templates ont été modifiés !")
print(f"📁 Originaux dans : {BACKUP_DIR}")
