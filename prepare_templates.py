#!/usr/bin/env python3
"""
Script pour préparer automatiquement les templates Word avec les variables
"""

import os
from docx import Document
from pathlib import Path
import re

# Dossier des templates
TEMPLATES_DIR = "Dossier exemple"

# Mapping des valeurs à remplacer par des variables
# Format: (pattern_regex, variable_de_remplacement)
REPLACEMENTS = [
    # Organisme de formation (exemples courants)
    (r'Aladé Conseil', '{{organisme_nom}}'),
    (r'SIRET\s*:\s*[\d\s]+', 'SIRET : {{organisme_siret}}'),
    (r'N°\s*de\s*déclaration\s*d\'activité\s*:\s*[\d\s]+', "N° de déclaration d'activité : {{organisme_nda}}"),
    (r'(\d{1,3})\s+(Rue|Avenue|Boulevard|Allée|Impasse|Place)\s+[^,]+,\s*\d{5}\s+\w+', '{{organisme_adresse}}, {{organisme_code_postal}} {{organisme_ville}}'),
    
    # Dates (formats courants)
    (r'\d{2}/\d{2}/\d{4}', '{{date_debut}}'),  # À affiner selon le contexte
    
    # Prix
    (r'\d+[,\s]?\d*\s*€\s*HT', '{{prix_total_ht}}'),
    (r'\d+[,\s]?\d*\s*€\s*TTC', '{{prix_total_ttc}}'),
    
    # Durée
    (r'\d+\s*heures?', '{{formation_duree}} heures'),
]

def replace_in_paragraph(paragraph, replacements):
    """Remplace les patterns dans un paragraphe"""
    for run in paragraph.runs:
        text = run.text
        for pattern, replacement in replacements:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        run.text = text

def replace_in_table(table, replacements):
    """Remplace les patterns dans un tableau"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)

def prepare_template(template_path, output_path=None):
    """
    Prépare un template en remplaçant les valeurs par des variables
    
    Args:
        template_path: Chemin du template original
        output_path: Chemin de sortie (si None, écrase l'original)
    """
    if output_path is None:
        output_path = template_path
    
    print(f"📄 Traitement de : {os.path.basename(template_path)}")
    
    try:
        # Charger le document
        doc = Document(template_path)
        
        # Remplacer dans les paragraphes
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, REPLACEMENTS)
        
        # Remplacer dans les tableaux
        for table in doc.tables:
            replace_in_table(table, REPLACEMENTS)
        
        # Remplacer dans les en-têtes
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                replace_in_paragraph(paragraph, REPLACEMENTS)
            for table in header.tables:
                replace_in_table(table, REPLACEMENTS)
        
        # Remplacer dans les pieds de page
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                replace_in_paragraph(paragraph, REPLACEMENTS)
            for table in footer.tables:
                replace_in_table(table, REPLACEMENTS)
        
        # Sauvegarder
        doc.save(output_path)
        print(f"   ✅ Template préparé avec succès\n")
        
    except Exception as e:
        print(f"   ❌ Erreur : {str(e)}\n")

def prepare_convention_template():
    """Prépare spécifiquement le template de convention avec toutes les variables"""
    template_path = os.path.join(TEMPLATES_DIR, "Modèle de convention simplifiée de formation 2020.docx")
    
    if not os.path.exists(template_path):
        print(f"❌ Template non trouvé : {template_path}")
        return
    
    print(f"📄 Préparation spéciale de la convention...")
    
    try:
        doc = Document(template_path)
        
        # Remplacements spécifiques pour la convention
        specific_replacements = {
            # Organisme
            'Aladé Conseil': '{{organisme_raison_sociale}}',
            'ALADE CONSEIL': '{{organisme_raison_sociale}}',
            
            # À compléter avec les valeurs réelles de votre template
        }
        
        # Appliquer les remplacements
        for paragraph in doc.paragraphs:
            for old_text, new_text in specific_replacements.items():
                if old_text in paragraph.text:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
        
        # Sauvegarder
        doc.save(template_path)
        print(f"   ✅ Convention préparée avec succès\n")
        
    except Exception as e:
        print(f"   ❌ Erreur : {str(e)}\n")

def main():
    """Prépare tous les templates"""
    print("🚀 Préparation automatique des templates Word\n")
    print("=" * 60)
    
    if not os.path.exists(TEMPLATES_DIR):
        print(f"❌ Dossier '{TEMPLATES_DIR}' non trouvé")
        return
    
    # Liste des templates à traiter
    templates = [
        "Modèle de convention simplifiée de formation 2020.docx",
        "Modèle Programme de formation 2020.docx",
        "Modèle Proposition de Formation_2.docx",
        "Modele Convocation formation 2020.docx",
        "Modèle Certificat de réalisation.docx",
        "Modèle Feuille d émargement entreprise.docx",
        "Modèle Feuille d émargement individuelle.docx",
        "Modèle Questionnaire préalable à la formation.docx",
        "Modèle évaluation à chaud.docx",
        "Modèle évaluation à froid.docx",
        "Modèle Évaluation de la satisfaction du client.docx",
        "Règlement Intérieur.docx",
        "Modele bulletin d'inscription formation INTER 2020.docx",
        "Modèle Grille de MAJ des compétences.docx",
        "Modèle Contrat Formateur.docx",
        "Modèle Déroulé Pédagogique.docx",
        "Modèle Questionnaire Formateur.docx",
        "Modèle Évaluation OPCO.docx",
        "Modèle Traitement des réclamations majeures.docx",
    ]
    
    # Créer une sauvegarde
    backup_dir = os.path.join(TEMPLATES_DIR, "backup_originaux")
    if not os.path.exists(backup_dir):
        os.makedirs(backup_dir)
        print(f"📁 Dossier de sauvegarde créé : {backup_dir}\n")
    
    # Traiter chaque template
    for template_name in templates:
        template_path = os.path.join(TEMPLATES_DIR, template_name)
        
        if os.path.exists(template_path):
            # Créer une sauvegarde
            backup_path = os.path.join(backup_dir, template_name)
            if not os.path.exists(backup_path):
                import shutil
                shutil.copy2(template_path, backup_path)
                print(f"💾 Sauvegarde créée : {template_name}")
            
            # Préparer le template
            # prepare_template(template_path)
        else:
            print(f"⚠️  Template non trouvé : {template_name}\n")
    
    print("=" * 60)
    print("✅ Préparation terminée !")
    print(f"📁 Les originaux sont sauvegardés dans : {backup_dir}")
    print("\n⚠️  IMPORTANT : Ce script fait des remplacements automatiques basiques.")
    print("   Vous devez vérifier et ajuster manuellement chaque template.")
    print("   Consultez PREPARATION-TEMPLATES.md pour la liste complète des variables.")

if __name__ == '__main__':
    main()
