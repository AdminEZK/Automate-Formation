"""
Service de Génération de Documents à partir de Templates Word
==============================================================
Utilise les modèles Word existants et remplace les variables
Date: 27 octobre 2025
Version: 1.0
"""

import os
from datetime import datetime
from docx import Document
from typing import Dict, List, Optional
import re
from pathlib import Path


class TemplateDocumentGenerator:
    """Générateur de documents à partir de templates Word"""
    
    def __init__(self, supabase_client, templates_dir: str = None):
        """
        Initialise le générateur de documents
        
        Args:
            supabase_client: Client Supabase pour accéder aux données
            templates_dir: Chemin vers le dossier des templates
        """
        self.supabase = supabase_client
        self.templates_dir = templates_dir or "Dossier exemple"
        self.output_dir = "generated_documents"
        self.ensure_output_dir()
        
    def ensure_output_dir(self):
        """Crée le dossier de sortie s'il n'existe pas"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    # ============================================
    # MÉTHODES UTILITAIRES
    # ============================================
    
    def get_session_data(self, session_id: str) -> Dict:
        """Récupère toutes les données d'une session"""
        response = self.supabase.table('vue_sessions_complete').select('*').eq('id', session_id).execute()
        if not response.data:
            raise ValueError(f"Session {session_id} non trouvée")
        return response.data[0]
    
    def get_organisme_data(self) -> Dict:
        """Récupère les données de l'organisme de formation"""
        response = self.supabase.table('organisme_formation').select('*').limit(1).execute()
        if not response.data:
            raise ValueError("Aucun organisme de formation configuré")
        return response.data[0]
    
    def get_participants(self, session_id: str) -> List[Dict]:
        """Récupère la liste des participants d'une session"""
        response = self.supabase.table('participants').select('*').eq('session_formation_id', session_id).execute()
        return response.data
    
    def get_formateur_data(self, formateur_id: str) -> Dict:
        """Récupère les données d'un formateur"""
        if not formateur_id:
            return {}
        response = self.supabase.table('formateurs').select('*').eq('id', formateur_id).execute()
        if not response.data:
            return {}
        return response.data[0]
    
    def format_date(self, date_str: str, format: str = "%d/%m/%Y") -> str:
        """Formate une date au format français"""
        if not date_str:
            return ""
        try:
            date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            return date_obj.strftime(format)
        except:
            return date_str
    
    def format_prix(self, prix: float) -> str:
        """Formate un prix en euros"""
        if prix is None:
            return "0,00 €"
        return f"{prix:,.2f} €".replace(',', ' ').replace('.', ',')
    
    def prepare_variables(self, session_id: str, participant_id: str = None) -> Dict:
        """
        Prépare toutes les variables pour le remplacement dans les templates
        
        Args:
            session_id: ID de la session
            participant_id: ID du participant (optionnel)
            
        Returns:
            Dictionnaire de variables {nom_variable: valeur}
        """
        session = self.get_session_data(session_id)
        organisme = self.get_organisme_data()
        
        # Variables de base
        variables = {
            # Organisme
            '{{organisme_nom}}': organisme.get('nom', ''),
            '{{organisme_raison_sociale}}': organisme.get('raison_sociale', ''),
            '{{organisme_siret}}': organisme.get('siret', ''),
            '{{organisme_nda}}': organisme.get('numero_declaration_activite', ''),
            '{{organisme_adresse}}': organisme.get('adresse', ''),
            '{{organisme_code_postal}}': organisme.get('code_postal', ''),
            '{{organisme_ville}}': organisme.get('ville', ''),
            '{{organisme_email}}': organisme.get('email', ''),
            '{{organisme_telephone}}': organisme.get('telephone', ''),
            '{{organisme_representant_nom}}': f"{organisme.get('representant_legal_prenom', '')} {organisme.get('representant_legal_nom', '')}",
            '{{organisme_representant_fonction}}': organisme.get('representant_legal_fonction', ''),
            
            # Entreprise cliente
            '{{entreprise_nom}}': session.get('entreprise_nom', ''),
            '{{entreprise_siret}}': session.get('entreprise_siret', ''),
            '{{entreprise_adresse}}': session.get('entreprise_adresse', ''),
            '{{entreprise_code_postal}}': session.get('entreprise_code_postal', ''),
            '{{entreprise_ville}}': session.get('entreprise_ville', ''),
            '{{entreprise_email}}': session.get('entreprise_email', ''),
            '{{entreprise_telephone}}': session.get('entreprise_telephone', ''),
            '{{entreprise_secteur_activite}}': session.get('entreprise_secteur_activite', ''),
            '{{entreprise_representant_nom}}': session.get('entreprise_representant_nom', ''),
            '{{entreprise_representant_fonction}}': session.get('entreprise_representant_fonction', ''),
            
            # Formation
            '{{formation_titre}}': session.get('formation_titre', ''),
            '{{formation_description}}': session.get('formation_description', ''),
            '{{formation_duree}}': str(session.get('formation_duree', '')),
            '{{formation_duree_jours}}': str(round(session.get('formation_duree', 0) / 7, 2)) if session.get('formation_duree') else '0',
            '{{formation_nature_action}}': session.get('formation_nature_action', 'Actions d\'acquisition, d\'entretien ou de perfectionnement des connaissances'),
            '{{formation_objectifs}}': session.get('formation_objectifs', ''),
            '{{formation_objectifs_pedagogiques}}': session.get('formation_objectifs_pedagogiques', ''),
            '{{formation_prerequis}}': session.get('formation_prerequis', 'Aucun'),
            '{{formation_programme}}': session.get('formation_programme', ''),
            '{{formation_methodes}}': session.get('formation_methodes', 'Méthodes actives et participatives'),
            '{{formation_moyens}}': session.get('formation_moyens', 'Supports pédagogiques, exercices pratiques'),
            '{{formation_public_vise}}': session.get('formation_public_vise', 'Tout public'),
            '{{formation_modalite}}': session.get('modalite', 'Présentiel'),
            
            # Session
            '{{session_id}}': session.get('id', ''),
            '{{session_numero}}': session.get('numero_session', ''),
            '{{date_debut}}': self.format_date(session.get('date_debut')),
            '{{date_fin}}': self.format_date(session.get('date_fin')),
            '{{lieu}}': session.get('lieu', ''),
            '{{nombre_participants}}': str(session.get('nombre_participants', '')),
            '{{horaire_debut}}': session.get('horaire_debut', '09:00'),
            '{{horaire_fin}}': session.get('horaire_fin', '17:00'),
            
            # Prix
            '{{prix_unitaire_ht}}': self.format_prix(session.get('formation_prix_ht', 0)),
            '{{prix_total_ht}}': self.format_prix(session.get('prix_total_ht', 0)),
            '{{prix_total_ttc}}': self.format_prix(session.get('prix_total_ht', 0) * 1.20),
            '{{prix_unitaire_ttc}}': self.format_prix(session.get('formation_prix_ht', 0) * 1.20),
            '{{tva}}': self.format_prix(session.get('prix_total_ht', 0) * 0.20),
            '{{frais_deplacement}}': self.format_prix(session.get('frais_deplacement', 0)),
            '{{prix_total_avec_frais}}': self.format_prix((session.get('prix_total_ht', 0) * 1.20) + session.get('frais_deplacement', 0)),
            
            # Dates système
            '{{date_aujourd_hui}}': datetime.now().strftime("%d/%m/%Y"),
            '{{date_proposition}}': datetime.now().strftime("%d/%m/%Y"),
            '{{date_entrevue}}': session.get('date_entrevue', ''),
            '{{annee}}': datetime.now().strftime("%Y"),
        }
        
        # Formateur si disponible
        if session.get('formateur_id'):
            formateur = self.get_formateur_data(session['formateur_id'])
            variables.update({
                '{{formateur_nom}}': formateur.get('nom', ''),
                '{{formateur_prenom}}': formateur.get('prenom', ''),
                '{{formateur_nom_complet}}': f"{formateur.get('prenom', '')} {formateur.get('nom', '')}",
                '{{formateur_email}}': formateur.get('email', ''),
                '{{formateur_telephone}}': formateur.get('telephone', ''),
            })
        else:
            variables.update({
                '{{formateur_nom}}': 'À confirmer',
                '{{formateur_prenom}}': '',
                '{{formateur_nom_complet}}': 'À confirmer',
                '{{formateur_email}}': '',
                '{{formateur_telephone}}': '',
            })
        
        # Participant si fourni
        if participant_id:
            response = self.supabase.table('participants').select('*').eq('id', participant_id).execute()
            if response.data:
                participant = response.data[0]
                variables.update({
                    '{{participant_nom}}': participant.get('nom', ''),
                    '{{participant_prenom}}': participant.get('prenom', ''),
                    '{{participant_nom_complet}}': f"{participant.get('prenom', '')} {participant.get('nom', '')}",
                    '{{participant_email}}': participant.get('email', ''),
                    '{{participant_telephone}}': participant.get('telephone', ''),
                    '{{participant_fonction}}': participant.get('fonction', ''),
                })
        
        return variables
    
    def replace_text_in_paragraph(self, paragraph, variables: Dict):
        """Remplace les variables dans un paragraphe"""
        for run in paragraph.runs:
            for var, value in variables.items():
                if var in run.text:
                    run.text = run.text.replace(var, str(value))
    
    def replace_text_in_table(self, table, variables: Dict):
        """Remplace les variables dans un tableau"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.replace_text_in_paragraph(paragraph, variables)
    
    def replace_variables_in_document(self, doc: Document, variables: Dict):
        """
        Remplace toutes les variables dans un document Word
        
        Args:
            doc: Document Word
            variables: Dictionnaire des variables à remplacer
        """
        # Remplacer dans les paragraphes
        for paragraph in doc.paragraphs:
            self.replace_text_in_paragraph(paragraph, variables)
        
        # Remplacer dans les tableaux
        for table in doc.tables:
            self.replace_text_in_table(table, variables)
        
        # Remplacer dans les en-têtes
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self.replace_text_in_paragraph(paragraph, variables)
            for table in header.tables:
                self.replace_text_in_table(table, variables)
        
        # Remplacer dans les pieds de page
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                self.replace_text_in_paragraph(paragraph, variables)
            for table in footer.tables:
                self.replace_text_in_table(table, variables)
    
    # ============================================
    # GÉNÉRATION DE DOCUMENTS
    # ============================================
    
    def generer_document_from_template(
        self, 
        template_name: str, 
        session_id: str, 
        participant_id: str = None,
        output_filename: str = None
    ) -> str:
        """
        Génère un document à partir d'un template
        
        Args:
            template_name: Nom du fichier template (ex: "Modèle de convention simplifiée de formation 2020.docx")
            session_id: ID de la session
            participant_id: ID du participant (optionnel)
            output_filename: Nom du fichier de sortie (optionnel)
            
        Returns:
            Chemin du fichier généré
        """
        # Chemin du template
        template_path = os.path.join(self.templates_dir, template_name)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template non trouvé: {template_path}")
        
        # Charger le template
        doc = Document(template_path)
        
        # Préparer les variables
        variables = self.prepare_variables(session_id, participant_id)
        
        # Remplacer les variables
        self.replace_variables_in_document(doc, variables)
        
        # Nom du fichier de sortie
        if not output_filename:
            base_name = os.path.splitext(template_name)[0]
            suffix = f"_{participant_id}" if participant_id else ""
            output_filename = f"{base_name}_{session_id}{suffix}.docx"
        
        output_path = os.path.join(self.output_dir, output_filename)
        
        # Sauvegarder le document
        doc.save(output_path)
        
        return output_path
    
    # ============================================
    # MÉTHODES SPÉCIFIQUES PAR TYPE DE DOCUMENT
    # ============================================
    
    def generer_convention(self, session_id: str) -> str:
        """Génère la convention de formation"""
        return self.generer_document_from_template(
            "Modèle de convention simplifiée de formation 2020.docx",
            session_id,
            output_filename=f"convention_formation_{session_id}.docx"
        )
    
    def generer_programme(self, session_id: str) -> str:
        """Génère le programme de formation"""
        return self.generer_document_from_template(
            "Modèle Programme de formation 2020.docx",
            session_id,
            output_filename=f"programme_formation_{session_id}.docx"
        )
    
    def generer_proposition(self, session_id: str) -> str:
        """Génère la proposition de formation"""
        return self.generer_document_from_template(
            "Modèle Proposition de Formation_2.docx",
            session_id,
            output_filename=f"proposition_formation_{session_id}.docx"
        )
    
    def generer_convocation(self, session_id: str, participant_id: str) -> str:
        """Génère une convocation pour un participant"""
        return self.generer_document_from_template(
            "Modele Convocation formation 2020.docx",
            session_id,
            participant_id,
            output_filename=f"convocation_{participant_id}.docx"
        )
    
    def generer_certificat(self, session_id: str, participant_id: str) -> str:
        """Génère un certificat de réalisation pour un participant"""
        return self.generer_document_from_template(
            "Modèle Certificat de réalisation.docx",
            session_id,
            participant_id,
            output_filename=f"certificat_{participant_id}.docx"
        )
    
    def generer_feuille_emargement_entreprise(self, session_id: str) -> str:
        """Génère la feuille d'émargement entreprise"""
        return self.generer_document_from_template(
            "Modèle Feuille d émargement entreprise.docx",
            session_id,
            output_filename=f"emargement_entreprise_{session_id}.docx"
        )
    
    def generer_feuille_emargement_individuelle(self, session_id: str, participant_id: str) -> str:
        """Génère la feuille d'émargement individuelle pour un participant"""
        return self.generer_document_from_template(
            "Modèle Feuille d émargement individuelle.docx",
            session_id,
            participant_id,
            output_filename=f"emargement_individuel_{participant_id}.docx"
        )
    
    def generer_questionnaire_prealable(self, session_id: str, participant_id: str) -> str:
        """Génère le questionnaire préalable pour un participant"""
        return self.generer_document_from_template(
            "Modèle Questionnaire préalable à la formation.docx",
            session_id,
            participant_id,
            output_filename=f"questionnaire_prealable_{participant_id}.docx"
        )
    
    def generer_evaluation_chaud(self, session_id: str, participant_id: str) -> str:
        """Génère l'évaluation à chaud pour un participant"""
        return self.generer_document_from_template(
            "Modèle évaluation à chaud.docx",
            session_id,
            participant_id,
            output_filename=f"evaluation_chaud_{participant_id}.docx"
        )
    
    def generer_evaluation_froid(self, session_id: str, participant_id: str) -> str:
        """Génère l'évaluation à froid pour un participant"""
        return self.generer_document_from_template(
            "Modèle évaluation à froid.docx",
            session_id,
            participant_id,
            output_filename=f"evaluation_froid_{participant_id}.docx"
        )
    
    def generer_evaluation_client(self, session_id: str) -> str:
        """Génère l'évaluation de satisfaction client"""
        return self.generer_document_from_template(
            "Modèle Évaluation de la satisfaction du client.docx",
            session_id,
            output_filename=f"evaluation_client_{session_id}.docx"
        )
    
    def generer_reglement_interieur(self, session_id: str) -> str:
        """Génère le règlement intérieur"""
        return self.generer_document_from_template(
            "Règlement Intérieur.docx",
            session_id,
            output_filename=f"reglement_interieur_{session_id}.docx"
        )
    
    def generer_bulletin_inscription(self, session_id: str, participant_id: str) -> str:
        """Génère le bulletin d'inscription pour un participant"""
        return self.generer_document_from_template(
            "Modele bulletin d'inscription formation INTER 2020.docx",
            session_id,
            participant_id,
            output_filename=f"bulletin_inscription_{participant_id}.docx"
        )
    
    def generer_grille_competences(self, session_id: str, participant_id: str) -> str:
        """Génère la grille de mise à jour des compétences"""
        return self.generer_document_from_template(
            "Modèle Grille de MAJ des compétences.docx",
            session_id,
            participant_id,
            output_filename=f"grille_competences_{participant_id}.docx"
        )
    
    def generer_contrat_formateur(self, session_id: str) -> str:
        """Génère le contrat formateur"""
        return self.generer_document_from_template(
            "Modèle Contrat Formateur.docx",
            session_id,
            output_filename=f"contrat_formateur_{session_id}.docx"
        )
    
    def generer_deroule_pedagogique(self, session_id: str) -> str:
        """Génère le déroulé pédagogique"""
        return self.generer_document_from_template(
            "Modèle Déroulé Pédagogique.docx",
            session_id,
            output_filename=f"deroule_pedagogique_{session_id}.docx"
        )
    
    def generer_questionnaire_formateur(self, session_id: str) -> str:
        """Génère le questionnaire formateur"""
        return self.generer_document_from_template(
            "Modèle Questionnaire Formateur.docx",
            session_id,
            output_filename=f"questionnaire_formateur_{session_id}.docx"
        )
    
    def generer_evaluation_opco(self, session_id: str) -> str:
        """Génère l'évaluation OPCO"""
        return self.generer_document_from_template(
            "Modèle Évaluation OPCO.docx",
            session_id,
            output_filename=f"evaluation_opco_{session_id}.docx"
        )
    
    def generer_traitement_reclamations(self, session_id: str) -> str:
        """Génère le document de traitement des réclamations"""
        return self.generer_document_from_template(
            "Modèle Traitement des réclamations majeures.docx",
            session_id,
            output_filename=f"traitement_reclamations_{session_id}.docx"
        )
    
    # ============================================
    # GÉNÉRATION COMPLÈTE
    # ============================================
    
    def generer_tous_documents_session(self, session_id: str) -> Dict[str, List[str]]:
        """
        Génère tous les documents pour une session
        
        Returns:
            Dict avec les chemins des fichiers générés par type
        """
        documents_generes = {
            'proposition': [],
            'convention': [],
            'programme': [],
            'convocations': [],
            'certificats': [],
            'emargements': [],
            'evaluations': [],
            'autres': []
        }
        
        try:
            # Documents session
            documents_generes['proposition'].append(self.generer_proposition(session_id))
            documents_generes['convention'].append(self.generer_convention(session_id))
            documents_generes['programme'].append(self.generer_programme(session_id))
            documents_generes['autres'].append(self.generer_reglement_interieur(session_id))
            documents_generes['emargements'].append(self.generer_feuille_emargement_entreprise(session_id))
            
            # Documents par participant
            participants = self.get_participants(session_id)
            for participant in participants:
                participant_id = participant['id']
                
                # Convocations
                documents_generes['convocations'].append(
                    self.generer_convocation(session_id, participant_id)
                )
                
                # Émargements individuels
                documents_generes['emargements'].append(
                    self.generer_feuille_emargement_individuelle(session_id, participant_id)
                )
                
                # Questionnaires et évaluations
                documents_generes['evaluations'].append(
                    self.generer_questionnaire_prealable(session_id, participant_id)
                )
                documents_generes['evaluations'].append(
                    self.generer_evaluation_chaud(session_id, participant_id)
                )
                documents_generes['evaluations'].append(
                    self.generer_evaluation_froid(session_id, participant_id)
                )
                
                # Certificats (à générer après la formation)
                # documents_generes['certificats'].append(
                #     self.generer_certificat(session_id, participant_id)
                # )
            
            # Évaluation client
            documents_generes['evaluations'].append(
                self.generer_evaluation_client(session_id)
            )
            
            return documents_generes
            
        except Exception as e:
            print(f"Erreur lors de la génération des documents: {str(e)}")
            raise
    
    # ============================================
    # MÉTHODES DE GÉNÉRATION PAR PHASE
    # ============================================
    
    def generer_phase_proposition(self, session_id):
        """
        PHASE 2 : Génère la proposition et le programme (J+1)
        """
        documents = []
        
        # Proposition de formation
        try:
            proposition = self.generer_proposition(session_id)
            documents.append({
                'type': 'proposition',
                'path': proposition,
                'name': 'Proposition de formation'
            })
        except Exception as e:
            print(f"Erreur génération proposition: {e}")
        
        # Programme de formation
        try:
            programme = self.generer_programme(session_id)
            documents.append({
                'type': 'programme',
                'path': programme,
                'name': 'Programme de formation'
            })
        except Exception as e:
            print(f"Erreur génération programme: {e}")
        
        return {
            'success': True,
            'documents': documents,
            'phase': 'proposition',
            'count': len(documents)
        }
    
    def generer_phase_preparation(self, session_id):
        """
        PHASE 4 : Génère les questionnaires préalables (J-7)
        """
        documents = []
        
        # Récupérer les participants
        participants = self.get_participants_data(session_id)
        
        for participant in participants:
            try:
                questionnaire = self.generer_questionnaire_prealable(session_id, participant['id'])
                documents.append({
                    'type': 'questionnaire_prealable',
                    'path': questionnaire,
                    'name': f"Questionnaire préalable - {participant['prenom']} {participant['nom']}"
                })
            except Exception as e:
                print(f"Erreur génération questionnaire pour {participant['nom']}: {e}")
        
        return {
            'success': True,
            'documents': documents,
            'phase': 'preparation',
            'count': len(documents)
        }
    
    def generer_phase_convocation(self, session_id):
        """
        PHASE 5 : Génère les convocations et feuilles d'émargement (J-4)
        """
        documents = []
        
        # Récupérer les participants
        participants = self.get_participants_data(session_id)
        
        # Convocations individuelles
        for participant in participants:
            try:
                convocation = self.generer_convocation(session_id, participant['id'])
                documents.append({
                    'type': 'convocation',
                    'path': convocation,
                    'name': f"Convocation - {participant['prenom']} {participant['nom']}"
                })
            except Exception as e:
                print(f"Erreur génération convocation pour {participant['nom']}: {e}")
        
        # Règlement intérieur (commun)
        try:
            reglement = self.generer_reglement_interieur(session_id)
            documents.append({
                'type': 'reglement_interieur',
                'path': reglement,
                'name': 'Règlement intérieur'
            })
        except Exception as e:
            print(f"Erreur génération règlement: {e}")
        
        # Feuille d'émargement entreprise
        try:
            emargement_entreprise = self.generer_feuille_emargement_entreprise(session_id)
            documents.append({
                'type': 'emargement_entreprise',
                'path': emargement_entreprise,
                'name': 'Feuille d\'émargement entreprise'
            })
        except Exception as e:
            print(f"Erreur génération émargement entreprise: {e}")
        
        # Feuilles d'émargement individuelles
        for participant in participants:
            try:
                emargement = self.generer_feuille_emargement_individuelle(session_id, participant['id'])
                documents.append({
                    'type': 'emargement_individuel',
                    'path': emargement,
                    'name': f"Feuille d\'émargement - {participant['prenom']} {participant['nom']}"
                })
            except Exception as e:
                print(f"Erreur génération émargement pour {participant['nom']}: {e}")
        
        return {
            'success': True,
            'documents': documents,
            'phase': 'convocation',
            'count': len(documents)
        }
    
    def generer_phase_evaluation_chaud(self, session_id):
        """
        PHASE 7 : Génère les évaluations à chaud (Fin formation)
        """
        documents = []
        
        # Récupérer les participants
        participants = self.get_participants_data(session_id)
        
        for participant in participants:
            try:
                evaluation = self.generer_evaluation_chaud(session_id, participant['id'])
                documents.append({
                    'type': 'evaluation_chaud',
                    'path': evaluation,
                    'name': f"Évaluation à chaud - {participant['prenom']} {participant['nom']}"
                })
            except Exception as e:
                print(f"Erreur génération évaluation chaud pour {participant['nom']}: {e}")
        
        return {
            'success': True,
            'documents': documents,
            'phase': 'evaluation_chaud',
            'count': len(documents)
        }
    
    def generer_phase_cloture(self, session_id):
        """
        PHASE 9 : Génère les certificats et l'évaluation client (J+2)
        """
        documents = []
        
        # Récupérer les participants
        participants = self.get_participants_data(session_id)
        
        # Certificats individuels
        for participant in participants:
            try:
                certificat = self.generer_certificat(session_id, participant['id'])
                documents.append({
                    'type': 'certificat',
                    'path': certificat,
                    'name': f"Certificat - {participant['prenom']} {participant['nom']}"
                })
            except Exception as e:
                print(f"Erreur génération certificat pour {participant['nom']}: {e}")
        
        # Évaluation satisfaction client
        try:
            evaluation_client = self.generer_evaluation_client(session_id)
            documents.append({
                'type': 'evaluation_client',
                'path': evaluation_client,
                'name': 'Évaluation satisfaction client'
            })
        except Exception as e:
            print(f"Erreur génération évaluation client: {e}")
        
        return {
            'success': True,
            'documents': documents,
            'phase': 'cloture',
            'count': len(documents)
        }
    
    def generer_phase_evaluation_froid(self, session_id):
        """
        PHASE 10 : Génère les évaluations à froid (J+60)
        """
        documents = []
        
        # Récupérer les participants
        participants = self.get_participants_data(session_id)
        
        for participant in participants:
            try:
                evaluation = self.generer_evaluation_froid(session_id, participant['id'])
                documents.append({
                    'type': 'evaluation_froid',
                    'path': evaluation,
                    'name': f"Évaluation à froid - {participant['prenom']} {participant['nom']}"
                })
            except Exception as e:
                print(f"Erreur génération évaluation froid pour {participant['nom']}: {e}")
        
        return {
            'success': True,
            'documents': documents,
            'phase': 'evaluation_froid',
            'count': len(documents)
        }

if __name__ == "__main__":
    from supabase import create_client
    import os
    from dotenv import load_dotenv
    import sys
    import json
    
    load_dotenv()
    
    # Configuration Supabase
    SUPABASE_URL = os.getenv('SUPABASE_URL')
    SUPABASE_KEY = os.getenv('SUPABASE_KEY')
    
    supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
    
    # Créer le générateur
    generator = TemplateDocumentGenerator(supabase)
    
    # Gestion des arguments en ligne de commande
    if len(sys.argv) > 1:
        method_name = sys.argv[1]
        session_id = sys.argv[2] if len(sys.argv) > 2 else None
        participant_id = sys.argv[3] if len(sys.argv) > 3 else None
        
        # Appeler la méthode demandée
        if hasattr(generator, method_name):
            method = getattr(generator, method_name)
            
            try:
                if participant_id:
                    result = method(session_id, participant_id)
                elif session_id:
                    result = method(session_id)
                else:
                    result = method()
                
                # Retourner le résultat en JSON
                print(json.dumps(result))
            except Exception as e:
                print(json.dumps({
                    'success': False,
                    'error': str(e)
                }))
        else:
            print(json.dumps({
                'success': False,
                'error': f'Méthode {method_name} non trouvée'
            }))
    else:
        # Mode test par défaut
        session_id = "votre-session-id"
        documents = generator.generer_tous_documents_session(session_id)
    
    print("Documents générés:")
    for type_doc, fichiers in documents.items():
        print(f"  {type_doc}: {len(fichiers)} fichier(s)")
        for fichier in fichiers:
            print(f"    - {fichier}")
